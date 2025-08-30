"""
Word to PDF Service
Handles converting Word documents (.docx, .doc) to PDF files with formatting preservation
"""
import os
import tempfile
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.lib.units import inch
from reportlab.lib.colors import black, blue, red, green
from reportlab.lib import colors
from flask import current_app
import mammoth
from utils import (
    generate_unique_id, create_temp_directory, 
    save_uploaded_files, cleanup_files, 
    secure_output_filename
)

class WordToPDFService:
    """Service class for handling Word to PDF conversion operations"""
    
    def __init__(self):
        self.upload_folder = current_app.config['UPLOAD_FOLDER']
    
    def convert_word_to_pdf(self, file, output_filename=None):
        """
        Convert Word document to PDF
        
        Args:
            file: Uploaded Word file object
            output_filename: Optional custom filename for output
            
        Returns:
            tuple: (success, result_data_or_error_message)
        """
        try:
            # Validate file
            is_valid, result = self._validate_word_file([file])
            if not is_valid:
                return False, result
            
            valid_file = result[0]
            
            # Generate unique ID and create temp directory
            unique_id = generate_unique_id()
            temp_dir = create_temp_directory(self.upload_folder, unique_id)
            
            # Save uploaded file temporarily
            temp_files = save_uploaded_files([valid_file], temp_dir)
            input_file = temp_files[0]
            
            # Prepare output filename
            if not output_filename:
                base_name = os.path.splitext(valid_file.filename)[0]
                output_filename = f"{base_name}.pdf"
            else:
                if not output_filename.endswith('.pdf'):
                    output_filename += '.pdf'
            
            output_filename = secure_output_filename(output_filename, default_name="word-to-pdf.pdf")
            
            # Convert Word to PDF
            success, result = self._convert_docx_to_pdf(input_file, temp_dir, output_filename)
            
            if success:
                # Clean up temporary files (but keep the PDF)
                cleanup_files(temp_files)
                
                return True, {
                    'unique_id': unique_id,
                    'filename': output_filename,
                    'message': 'Word document converted to PDF successfully!'
                }
            else:
                # Clean up all files on error
                cleanup_files(temp_files + [temp_dir])
                return False, result
                
        except Exception as e:
            return False, f'An error occurred during conversion: {str(e)}'
    
    def _validate_word_file(self, files):
        """
        Validate uploaded Word files
        
        Args:
            files: List of uploaded file objects
            
        Returns:
            tuple: (is_valid, result_or_error_message)
        """
        if not files or len(files) == 0:
            return False, 'No files provided'
        
        file = files[0]
        
        if not file or file.filename == '':
            return False, 'No file selected'
        
        # Check file extension
        allowed_extensions = ['.docx', '.doc']
        file_ext = os.path.splitext(file.filename.lower())[1]
        
        if file_ext not in allowed_extensions:
            return False, f'Invalid file type. Please upload a Word document (.docx or .doc)'
        
        # Check file size (16MB limit)
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        max_size = 16 * 1024 * 1024  # 16MB
        if file_size > max_size:
            return False, 'File size too large. Maximum allowed size is 16MB.'
        
        return True, [file]
    
    def _convert_docx_to_pdf(self, input_file, temp_dir, output_filename):
        """
        Internal method to convert DOCX to PDF with proper formatting preservation
        
        Args:
            input_file: Path to input Word file
            temp_dir: Temporary directory path
            output_filename: Name for the output PDF file
            
        Returns:
            tuple: (success, result_or_error)
        """
        try:
            # Read the Word document
            doc = Document(input_file)
            
            # Create PDF
            output_path = os.path.join(temp_dir, output_filename)
            pdf_doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Get base styles
            styles = getSampleStyleSheet()
            story = []
            
            # Process all document elements in order
            # First, collect all paragraphs and tables in order
            document_elements = []
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                document_elements.append(('paragraph', paragraph))
                
            # Process tables
            for table in doc.tables:
                document_elements.append(('table', table))
            
            # Sort by document order (this is a simplified approach)
            # In a real implementation, you'd need to track actual document order
            
            # Process paragraphs with preserved formatting
            for paragraph in doc.paragraphs:
                if paragraph.text.strip() or len(paragraph.text.strip()) == 0:
                    # Skip paragraphs that are part of tables
                    skip_paragraph = False
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if paragraph in cell.paragraphs:
                                    skip_paragraph = True
                                    break
                            if skip_paragraph:
                                break
                        if skip_paragraph:
                            break
                    
                    if skip_paragraph:
                        continue
                    
                    # Create custom style based on Word formatting
                    custom_style = self._create_paragraph_style(paragraph, styles)
                    
                    # Handle formatted text with runs
                    if paragraph.runs:
                        formatted_text = self._process_paragraph_runs(paragraph)
                        if formatted_text.strip():
                            story.append(Paragraph(formatted_text, custom_style))
                        else:
                            # Empty paragraph - add space
                            story.append(Spacer(1, 12))
                    elif paragraph.text.strip():
                        # Simple text paragraph
                        story.append(Paragraph(paragraph.text, custom_style))
                    else:
                        # Empty paragraph
                        story.append(Spacer(1, 12))
            
            # Process tables if any
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text if cell_text else " ")
                    table_data.append(row_data)
                
                if table_data:
                    # Create table
                    pdf_table = Table(table_data)
                    pdf_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(pdf_table)
                    story.append(Spacer(1, 12))
            
            # Handle empty document
            if not story:
                default_style = ParagraphStyle(
                    'Default',
                    parent=styles['Normal'],
                    fontSize=12,
                    alignment=TA_LEFT
                )
                story.append(Paragraph("This document appears to be empty or contains only images/tables.", default_style))
            
            # Build PDF
            pdf_doc.build(story)
            
            return True, output_path
            
        except Exception as e:
            current_app.logger.error(f'Error converting Word to PDF: {str(e)}')
            return False, f'Error converting Word to PDF: {str(e)}'
    
    def _create_paragraph_style(self, paragraph, base_styles):
        """
        Create a custom ReportLab style based on Word paragraph formatting
        
        Args:
            paragraph: Word paragraph object
            base_styles: ReportLab base styles
            
        Returns:
            ParagraphStyle: Custom style for this paragraph
        """
        # Determine alignment
        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: TA_LEFT,
            WD_PARAGRAPH_ALIGNMENT.CENTER: TA_CENTER,
            WD_PARAGRAPH_ALIGNMENT.RIGHT: TA_RIGHT,
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: TA_JUSTIFY,
        }
        
        alignment = alignment_map.get(paragraph.alignment, TA_LEFT)
        
        # Determine if this is a heading
        is_heading = paragraph.style.name.startswith('Heading') or self._is_heading(paragraph)
        
        # Base style selection
        if is_heading:
            if paragraph.style.name == 'Heading 1':
                base_style = base_styles['Heading1']
                font_size = 18
            elif paragraph.style.name == 'Heading 2':
                base_style = base_styles['Heading2'] 
                font_size = 16
            elif paragraph.style.name == 'Heading 3':
                base_style = base_styles['Heading3']
                font_size = 14
            else:
                base_style = base_styles['Heading1']
                font_size = 16
        else:
            base_style = base_styles['Normal']
            font_size = 12
        
        # Check for bold in first run
        is_bold = False
        if paragraph.runs:
            is_bold = paragraph.runs[0].bold
        
        # Create custom style
        style_name = f'Custom_{hash(paragraph.text[:20])}'
        custom_style = ParagraphStyle(
            style_name,
            parent=base_style,
            fontSize=font_size,
            alignment=alignment,
            spaceAfter=12 if is_heading else 6,
            spaceBefore=6 if is_heading else 0,
            leftIndent=0,
            rightIndent=0,
        )
        
        # Apply additional formatting
        if is_bold and not is_heading:
            custom_style.fontName = 'Helvetica-Bold'
        
        return custom_style
    
    def _process_paragraph_runs(self, paragraph):
        """
        Process paragraph runs to preserve text formatting like bold, italic
        
        Args:
            paragraph: Word paragraph object
            
        Returns:
            str: HTML-formatted text for ReportLab
        """
        result = ""
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
                
            # Escape HTML special characters
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Apply formatting
            if run.bold:
                text = f"<b>{text}</b>"
            if run.italic:
                text = f"<i>{text}</i>"
            if run.underline:
                text = f"<u>{text}</u>"
            
            result += text
        
        return result
    
    def _is_heading(self, paragraph):
        """
        Determine if a paragraph is likely a heading
        
        Args:
            paragraph: Word paragraph object
            
        Returns:
            bool: True if likely a heading
        """
        # Check if the paragraph style indicates a heading
        if paragraph.style.name.startswith('Heading'):
            return True
        
        # Check for common heading characteristics
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if (first_run.bold and 
                len(paragraph.text.strip()) < 100 and 
                not paragraph.text.strip().endswith('.')):
                return True
        
        return False
    
    def get_file_path(self, unique_id, filename):
        """Get the full path to a processed file"""
        return os.path.join(self.upload_folder, unique_id, filename)
    
    def cleanup_processed_file(self, file_path, delay=None):
        """Clean up a processed file with optional delay"""
        cleanup_delay = delay or current_app.config.get('CLEANUP_DELAY', 10)
        
        temp_dir = os.path.dirname(file_path)
        cleanup_files([file_path, temp_dir], delay=cleanup_delay)